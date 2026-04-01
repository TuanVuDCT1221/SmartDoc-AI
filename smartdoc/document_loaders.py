from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile

import pdfplumber
from docx import Document as DocxDocument
from langchain_community.document_loaders import Docx2txtLoader
from pypdf import PdfReader
from pypdf.errors import PdfReadError


SUPPORTED_TYPES = {".pdf": "pdf", ".docx": "docx"}
CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")


@dataclass(slots=True)
class LoadedDocument:
    text: str
    source_name: str
    source_type: str


def normalize_text(text: str | bytes) -> str:
    if isinstance(text, bytes):
        text = text.decode("utf-8", errors="replace")

    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = CONTROL_CHARS_RE.sub(" ", text)
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _resolve_source_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[suffix]
    raise ValueError("Unsupported file type. Supported file types are: .pdf, .docx")


def load_document(file_path: str | Path) -> LoadedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise ValueError(f"Expected a file path, got: {path}")

    source_type = _resolve_source_type(path)

    try:
        if source_type == "pdf":
            text = extract_pdf_text(path)
        else:
            text = extract_docx_text(path)
    except Exception as exc:  
        raise ValueError(f"Failed to parse '{path.name}' as {source_type.upper()}: {exc}") from exc

    text = normalize_text(text)

    if not text.strip():
        raise ValueError("The uploaded file did not produce readable text.")

    return LoadedDocument(
        text=text,
        source_name=path.name,
        source_type=source_type,
    )


def extract_pdf_text(path: Path) -> str:
    errors: list[str] = []

    try:
        text = _extract_pdf_text_with_pypdf(path)
        if text:
            return text
    except (PdfReadError, OSError, ValueError) as exc:
        errors.append(f"pypdf: {exc}")

    try:
        text = _extract_pdf_text_with_pdfplumber(path)
        if text:
            return text
    except Exception as exc:  
        errors.append(f"pdfplumber: {exc}")

    if errors:
        raise RuntimeError("; ".join(errors))
    raise ValueError("No readable text was extracted from the PDF file.")


def _extract_pdf_text_with_pypdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = normalize_text(page.extract_text() or "")
        if page_text:
            pages.append(f"[Page {page_number}]\n{page_text}")

    if not pages:
        raise ValueError("pypdf could not extract readable text.")

    return "\n\n".join(pages)


def _extract_pdf_text_with_pdfplumber(path: Path) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = normalize_text(page.extract_text() or "")
            if page_text:
                pages.append(f"[Page {page_number}]\n{page_text}")

    if not pages:
        raise ValueError("pdfplumber could not extract readable text.")

    return "\n\n".join(pages)


def extract_docx_text(path: Path) -> str:
    errors: list[str] = []

    try:
        text = _extract_docx_text_with_python_docx(path)
        if text:
            return text
    except (BadZipFile, OSError, ValueError) as exc:
        errors.append(f"python-docx: {exc}")
    except Exception as exc:  
        errors.append(f"python-docx: {exc}")

    try:
        text = _extract_docx_text_with_loader(path)
        if text:
            return text
    except Exception as exc: 
        errors.append(f"Docx2txtLoader: {exc}")

    if errors:
        raise RuntimeError("; ".join(errors))
    raise ValueError("No readable text was extracted from the DOCX file.")


def _extract_docx_text_with_python_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        heading_level = _heading_level_from_style(style_name)
        if heading_level:
            parts.append(f"{'#' * heading_level} {text}")
        else:
            parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        markdown_rows = _table_to_markdown_rows(table)
        if markdown_rows:
            parts.append(f"[Table {table_index}]")
            parts.extend(markdown_rows)

    return "\n\n".join(parts)


def _extract_docx_text_with_loader(path: Path) -> str:
    loader = Docx2txtLoader(str(path))
    documents = loader.load()
    parts: list[str] = []

    for doc in documents:
        page_content = normalize_text(doc.page_content)
        if page_content:
            parts.append(page_content)

    if not parts:
        raise ValueError("Docx2txtLoader returned empty content.")

    return "\n\n".join(parts)


def _heading_level_from_style(style_name: str) -> int | None:
    match = re.search(r"heading\s*(\d+)", style_name.lower())
    if not match:
        return None
    level = int(match.group(1))
    return max(1, min(level, 6))


def _table_to_markdown_rows(table: object) -> list[str]:
    rows: list[list[str]] = []

    for row in table.rows:  # type: ignore[attr-defined]
        cells = [normalize_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append([cell if cell else "-" for cell in cells])

    if not rows:
        return []

    max_columns = max(len(row) for row in rows)

    normalized_rows: list[list[str]] = []
    for row in rows:
        if len(row) < max_columns:
            row = row + ["-"] * (max_columns - len(row))
        normalized_rows.append(row)

    markdown_rows: list[str] = [" | ".join(normalized_rows[0])]
    if len(normalized_rows) > 1:
        markdown_rows.append(" | ".join(["---"] * max_columns))
        markdown_rows.extend(" | ".join(row) for row in normalized_rows[1:])

    return markdown_rows
