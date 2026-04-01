from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document as DocxDocument

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from smartdoc.document_loaders import load_document


def create_sample_docx(target_path: Path) -> Path:
    target_path.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()
    doc.add_heading("SmartDoc QA Validation Document", level=1)
    doc.add_paragraph(
        "This sample is generated to validate DOCX extraction quality for headings, tables, and long paragraphs."
    )

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "SmartDoc should preserve heading markers clearly so the retriever can keep section context stable and reliable."
    )

    doc.add_heading("Long Context Section", level=2)
    long_text = " ".join(
        [
            "This is an intentionally long paragraph used for extraction stress testing."
            " It includes repeated clauses to emulate legal, policy, or technical documents."
        ]
        * 30
    )
    doc.add_paragraph(long_text)

    doc.add_heading("Comparison Table", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Metric"
    headers[1].text = "Current"
    headers[2].text = "Target"

    rows = [
        ("Heading coverage", "2", "3+"),
        ("Table extraction", "basic", "markdown rows"),
        ("Long paragraph handling", "partial", "stable"),
    ]
    for metric, current, target in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = current
        row_cells[2].text = target

    doc.add_heading("Implementation Notes", level=3)
    doc.add_paragraph(
        "When parsing production files, fallback loaders help with uncommon formatting edge cases."
    )

    doc.save(target_path)
    return target_path


def summarize_extraction(docx_path: Path) -> int:
    loaded = load_document(docx_path)
    text = loaded.text

    heading_count = len(re.findall(r"^#{1,6}\s", text, flags=re.MULTILINE))
    table_count = text.count("[Table ")
    paragraphs = [segment.strip() for segment in text.split("\n\n") if segment.strip()]
    longest_paragraph = max((len(paragraph) for paragraph in paragraphs), default=0)

    print(f"Document: {loaded.source_name} ({loaded.source_type})")
    print(f"Characters: {len(text)}")
    print(f"Headings detected: {heading_count}")
    print(f"Tables detected: {table_count}")
    print(f"Longest paragraph length: {longest_paragraph}")
    print("Preview:")
    print(text[:700])

    checks = [
        (heading_count >= 3, "Expected at least 3 headings."),
        (table_count >= 1, "Expected at least 1 table marker."),
        (longest_paragraph >= 400, "Expected long paragraph length >= 400 chars."),
    ]

    failed_messages = [message for passed, message in checks if not passed]
    if failed_messages:
        print("\nValidation failed:")
        for message in failed_messages:
            print(f"- {message}")
        return 1

    print("\nValidation passed.")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Validate DOCX extraction for headings, tables, and long paragraphs."
    )
    parser.add_argument(
        "--docx",
        type=Path,
        default=ROOT_DIR / "data" / "sample_complex.docx",
        help="Path to a DOCX file to validate.",
    )
    parser.add_argument(
        "--generate-sample",
        action="store_true",
        help="Generate a synthetic DOCX sample before running validation.",
    )
    args = parser.parse_args()

    target_docx = args.docx
    if args.generate_sample or not target_docx.exists():
        target_docx = create_sample_docx(target_docx)
        print(f"Sample DOCX generated at: {target_docx}")

    return summarize_extraction(target_docx)


if __name__ == "__main__":
    raise SystemExit(main())
